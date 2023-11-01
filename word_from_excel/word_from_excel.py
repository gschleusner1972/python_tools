import os
import csv
from docx import Document

BASE_DIR = "G:\\Git_2\\python_tools\\word_from_excel"

def get_table_title(table):
    """Retrieve the table title from the first cell."""
    return table.cell(0, 0).text.strip().lower()

def populate_table_in_doc(doc, name_value):
    """Populate a table in the Word document with data from a CSV file."""
    for table in doc.tables:
        table_title = get_table_title(table)
        if table_title:
            table_file = f"{table_title}_{name_value}.csv"
            file_path = os.path.join(BASE_DIR, table_file)
            
            if os.path.exists(file_path):
                with open(file_path, 'r', encoding="utf-8") as csv_file:
                    csv_reader = csv.reader(csv_file)
                    next(csv_reader)  # skip header row

                    # Delete all rows in the table except the header
                    while len(table.rows) > 1:
                        row = table.rows[-1]
                        row._element.get_or_add_tblPr().get_or_add_tblGrid().del_col()

                    # Populate the table with data from the CSV file
                    for row_data in csv_reader:
                        cells = table.add_row().cells
                        for i, value in enumerate(row_data):
                            cells[i].text = str(value)

def generate_word_from_csv():
    """Generate Word documents from CSV data."""
    with open(os.path.join(BASE_DIR, "data.csv"), 'r', encoding="utf-8") as csv_file:
        csv_reader = csv.reader(csv_file)
        headers = next(csv_reader)

        for row in csv_reader:
            doc = Document(os.path.join(BASE_DIR, "template.docx"))
            
            # Replace placeholders in the document
            for cell_value, header in zip(row, headers):
                for paragraph in doc.paragraphs:
                    paragraph.text = paragraph.text.replace(f"[[{header}]]", str(cell_value))
        
            # Get the 'name' value
            name_index = headers.index('name')
            name_value = row[name_index]

            if name_value:
                populate_table_in_doc(doc, name_value)
        
            output_dir = os.path.join(BASE_DIR, "output")
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)
            doc.save(os.path.join(output_dir, f"document_{name_value}.docx"))

if __name__ == "__main__":
    generate_word_from_csv()
